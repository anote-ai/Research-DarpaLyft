"""Generate polished DOCX and PDF documentation for the DARPA LIFT submission."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

TITLE = "Multirotor Flight Testing: Clarification for DARPA LIFT Challenge"
SUBTITLE = "How Test Flights Support Our Primary UAS Design"
META = [
    ("Submitted by", "Natan Vidra & Team Anote"),
    ("Organization", "Anote, Inc."),
    ("Date", "June 11, 2026"),
    ("Re", "DARPA LIFT Challenge — Response to Reviewer Query"),
]

SUMMARY = (
    "The multirotor flight tests visible in our uploaded footage are physics "
    "validation and calibration flights, not demonstrations of our final competition "
    "design. These flights are a deliberate step in our AI-driven design pipeline: "
    "we use real flight data to ground-truth the physics models that our optimizer "
    "relies on before committing designs to hardware. Our primary UAS design concept "
    "is detailed in the attached Concept Paper (Anote_Final_Concept_Paper.pdf)."
)

STAGES = [
    {
        "title": "Stage 1 — Physics Model Calibration (Multirotor Test Flights)",
        "body": (
            "We fly off-the-shelf multirotor platforms (quadcopters and hexacopters) "
            "to validate the core physics relationships our optimizer uses. The test "
            "vehicles are not the competition design — they are instrumented benches "
            "used to confirm that our simplified actuator disk model, structural "
            "equations, and feasibility constraints reflect real-world behavior before "
            "we run the optimizer at scale."
        ),
        "table": {
            "headers": ["Model Parameter", "Formula", "Validated Against"],
            "rows": [
                ["Thrust", "T = 10 × N_motors × D_prop² × throttle",
                 "Hover thrust at varying throttle"],
                ["Payload capacity", "payload = (T / 9.81) − total_mass",
                 "Known payload added incrementally"],
                ["Structural safety factor", "mean_UTS / 100 MPa",
                 "Load deflection of CF vs. Al frames"],
                ["Thrust-to-weight ratio", "T / (mass × 9.81) ≥ 1.5",
                 "Minimum throttle at various mass configs"],
            ],
        },
    },
    {
        "title": "Stage 2 — AI-Driven Design Optimization",
        "body": (
            "With validated physics models, we run our optimization pipeline "
            "(src/darpalyft/) over a large design space. The optimizer maximizes "
            "score = payload_kg / total_mass_kg subject to all physics and structural "
            "constraints, using Pareto front analysis to surface non-dominated designs."
        ),
        "bullets": [
            "Materials: carbon fiber, aluminum alloy, titanium, fiberglass, ABS plastic",
            "Motor count: 4, 6, or 8",
            "Propeller diameter: 0.2 – 0.5 m",
            "Battery capacity: 400 – 1,200 Wh",
            "Component masses: per-component continuous variables",
        ],
    },
    {
        "title": "Stage 3 — Primary Design Selection & Refinement",
        "body": (
            "The highest-scoring feasible designs from Stage 2 feed into our primary "
            "UAS design (detailed in the Concept Paper). The Concept Paper describes "
            "the specific airframe configuration, propulsion system, and structural "
            "approach we are proposing for the competition — informed by, but distinct "
            "from, the generic multirotor test platforms."
        ),
    },
]

RATIONALE = (
    "Testing the final design first would conflate calibration errors with design "
    "errors. By validating physics on known, off-the-shelf platforms — where ground "
    "truth is well-characterized — we ensure our optimizer is working from accurate "
    "physical models. Only after that validation do we trust the optimizer's output "
    "as a guide for the competition design. This approach also lets us iterate "
    "quickly: multirotor test platforms are cheap and fast to reconfigure, while "
    "the primary design involves more specialized manufacturing."
)

REF_TABLE = {
    "headers": ["Document", "Location", "Description"],
    "rows": [
        ["Concept Paper", "docs/Anote_Final_Concept_Paper.pdf",
         "Full primary UAS design concept for DARPA LIFT"],
        ["Physics model", "src/darpalyft/core.py",
         "Thrust, structural, and payload calculations"],
        ["Optimizer", "src/darpalyft/core.py — DesignOptimizer",
         "Design search and feasibility filtering"],
        ["Evaluation metrics", "src/darpalyft/evaluate.py",
         "Payload-per-weight scoring and Pareto analysis"],
        ["Optimization pipeline", "scripts/run_optimization.py",
         "End-to-end run script"],
    ],
}

CONTACT = "Natan Vidra — natan@anote.ai  |  Anote, Inc. — anote.ai"

# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

ANOTE_BLUE = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
ANOTE_DARK = RGBColor(0x11, 0x18, 0x27)   # #111827
GRAY      = RGBColor(0x6B, 0x72, 0x80)   # #6B7280


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = ANOTE_BLUE if level == 1 else ANOTE_DARK
    return p


def _add_table(doc, headers, rows):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], "1A56DB")
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F3F4F6" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            _set_cell_bg(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = ANOTE_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(SUBTITLE)
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Metadata table
    meta_table = doc.add_table(rows=len(META), cols=2)
    meta_table.style = "Table Grid"
    for i, (key, val) in enumerate(META):
        cells = meta_table.rows[i].cells
        cells[0].text = key
        cells[1].text = val
        _set_cell_bg(cells[0], "EFF6FF")
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = ANOTE_BLUE
        for para in cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # Summary
    _add_heading(doc, "Executive Summary")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

    # Pipeline stages
    _add_heading(doc, "How Multirotor Testing Fits the Design Pipeline")

    for stage in STAGES:
        _add_heading(doc, stage["title"], level=2)
        bp = doc.add_paragraph(stage["body"])
        bp.paragraph_format.space_after = Pt(4)
        for run in bp.runs:
            run.font.size = Pt(10)

        if "table" in stage:
            _add_table(doc, stage["table"]["headers"], stage["table"]["rows"])

        if "bullets" in stage:
            for bullet in stage["bullets"]:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(10)
            doc.add_paragraph()

    # Rationale
    _add_heading(doc, "Why Test Multirotors Instead of the Final Design Directly?")
    rp = doc.add_paragraph(RATIONALE)
    for run in rp.runs:
        run.font.size = Pt(10)

    # Reference documents
    _add_heading(doc, "Reference Documents")
    _add_table(doc, REF_TABLE["headers"], REF_TABLE["rows"])

    # Contact
    _add_heading(doc, "Contact")
    cp = doc.add_paragraph(CONTACT)
    for run in cp.runs:
        run.font.size = Pt(10)

    out_path = os.path.join(OUT_DIR, "Anote_DARPA_LIFT_Clarification.docx")
    doc.save(out_path)
    print(f"DOCX saved: {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

BLUE  = colors.HexColor("#1A56DB")
DARK  = colors.HexColor("#111827")
LGRAY = colors.HexColor("#F3F4F6")
MGRAY = colors.HexColor("#6B7280")
WHITE = colors.white


def build_pdf():
    out_path = os.path.join(OUT_DIR, "Anote_DARPA_LIFT_Clarification.pdf")
    doc = SimpleDocTemplate(
        out_path,
        pagesize=letter,
        leftMargin=1.25 * inch,
        rightMargin=1.25 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()

    s_title = ParagraphStyle(
        "DocTitle",
        fontSize=20,
        leading=26,
        textColor=BLUE,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    s_subtitle = ParagraphStyle(
        "DocSubtitle",
        fontSize=12,
        leading=16,
        textColor=MGRAY,
        fontName="Helvetica-Oblique",
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    s_h1 = ParagraphStyle(
        "H1",
        fontSize=13,
        leading=18,
        textColor=BLUE,
        fontName="Helvetica-Bold",
        spaceBefore=14,
        spaceAfter=4,
    )
    s_h2 = ParagraphStyle(
        "H2",
        fontSize=11,
        leading=15,
        textColor=DARK,
        fontName="Helvetica-Bold",
        spaceBefore=10,
        spaceAfter=3,
    )
    s_body = ParagraphStyle(
        "Body",
        fontSize=10,
        leading=14,
        textColor=DARK,
        fontName="Helvetica",
        spaceAfter=6,
        alignment=TA_JUSTIFY,
    )
    s_bullet = ParagraphStyle(
        "Bullet",
        fontSize=10,
        leading=14,
        textColor=DARK,
        fontName="Helvetica",
        leftIndent=18,
        bulletIndent=6,
        spaceAfter=3,
    )
    s_meta_key = ParagraphStyle(
        "MetaKey",
        fontSize=9,
        fontName="Helvetica-Bold",
        textColor=BLUE,
    )
    s_meta_val = ParagraphStyle(
        "MetaVal",
        fontSize=9,
        fontName="Helvetica",
        textColor=DARK,
    )
    s_contact = ParagraphStyle(
        "Contact",
        fontSize=10,
        leading=14,
        textColor=DARK,
        fontName="Helvetica",
        alignment=TA_CENTER,
    )

    def make_table(headers, rows, col_widths=None):
        data = [[Paragraph(f"<b>{h}</b>", ParagraphStyle(
            "TH", fontSize=9, fontName="Helvetica-Bold",
            textColor=WHITE, leading=12)
        ) for h in headers]]
        for r in rows:
            data.append([Paragraph(cell, ParagraphStyle(
                "TD", fontSize=9, fontName="Helvetica",
                textColor=DARK, leading=12)
            ) for cell in r])

        available = 6.5 * inch
        if col_widths is None:
            col_widths = [available / len(headers)] * len(headers)

        style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LGRAY, WHITE]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ])
        return Table(data, colWidths=col_widths, style=style, hAlign="LEFT")

    story = []

    # Title
    story.append(Paragraph(TITLE, s_title))
    story.append(Paragraph(SUBTITLE, s_subtitle))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLUE, spaceAfter=10))

    # Metadata
    meta_data = [[
        Paragraph(k, s_meta_key),
        Paragraph(v, s_meta_val),
    ] for k, v in META]
    meta_tbl = Table(meta_data, colWidths=[1.5 * inch, 5 * inch],
                     style=TableStyle([
                         ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#EFF6FF"), WHITE]),
                         ("TOPPADDING", (0, 0), (-1, -1), 4),
                         ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                         ("LEFTPADDING", (0, 0), (-1, -1), 6),
                     ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 14))

    # Summary
    story.append(Paragraph("Executive Summary", s_h1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"), spaceAfter=6))
    story.append(Paragraph(SUMMARY, s_body))

    # Pipeline
    story.append(Paragraph("How Multirotor Testing Fits the Design Pipeline", s_h1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"), spaceAfter=6))

    for stage in STAGES:
        block = []
        block.append(Paragraph(stage["title"], s_h2))
        block.append(Paragraph(stage["body"], s_body))

        if "table" in stage:
            t = stage["table"]
            cw = [2.0 * inch, 2.5 * inch, 2.0 * inch]
            block.append(make_table(t["headers"], t["rows"], col_widths=cw))
            block.append(Spacer(1, 8))

        if "bullets" in stage:
            for b in stage["bullets"]:
                block.append(Paragraph(f"•  {b}", s_bullet))
            block.append(Spacer(1, 6))

        story.append(KeepTogether(block))

    # Rationale
    story.append(Paragraph("Why Test Multirotors Instead of the Final Design Directly?", s_h1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"), spaceAfter=6))
    story.append(Paragraph(RATIONALE, s_body))

    # Reference documents
    story.append(Paragraph("Reference Documents", s_h1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D1D5DB"), spaceAfter=6))
    cw = [1.5 * inch, 2.5 * inch, 2.5 * inch]
    story.append(make_table(REF_TABLE["headers"], REF_TABLE["rows"], col_widths=cw))
    story.append(Spacer(1, 12))

    # Contact
    story.append(HRFlowable(width="100%", thickness=1, color=BLUE, spaceBefore=10, spaceAfter=8))
    story.append(Paragraph(CONTACT, s_contact))

    doc.build(story)
    print(f"PDF saved:  {out_path}")
    return out_path


if __name__ == "__main__":
    build_docx()
    build_pdf()
